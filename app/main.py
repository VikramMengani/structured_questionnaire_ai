from .rag import load_reference_docs
from .models import Question, Answer
from sqlalchemy.orm import Session
from docx import Document
from fastapi.responses import FileResponse
import numpy as np
from fastapi import FastAPI, UploadFile, File, Depends
from sqlalchemy.orm import Session
from .database import Base, engine, SessionLocal
from .models import User, Questionnaire, Question, Answer
from .auth import hash_password, verify_password, create_token
import shutil
import os

Base.metadata.create_all(bind=engine)

app = FastAPI()


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


@app.post("/signup")
def signup(email: str, password: str, db: Session = Depends(get_db)):
    user = User(email=email, password=hash_password(password))
    db.add(user)
    db.commit()
    return {"message": "User created"}


@app.post("/login")
def login(email: str, password: str, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == email).first()
    if not user or not verify_password(password, user.password):
        return {"error": "Invalid credentials"}

    token = create_token({"sub": user.email})
    return {"access_token": token}


@app.post("/upload_questionnaire")
def upload_questionnaire(file: UploadFile = File(...), db: Session = Depends(get_db)):

    file_path = f"uploads/{file.filename}"

    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # Create questionnaire record
    questionnaire = Questionnaire(file_path=file_path)
    db.add(questionnaire)
    db.commit()
    db.refresh(questionnaire)

    # Read file and split into questions
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()

    # Split by line (each line = one question)
    lines = [line.strip() for line in content.split("\n") if line.strip()]

    for line in lines:
        question = Question(
            questionnaire_id=questionnaire.id,
            text=line
        )
        db.add(question)

    db.commit()

    return {
        "message": "Uploaded and questions saved",
        "questionnaire_id": questionnaire.id,
        "total_questions": len(lines)
    }

@app.post("/generate_answers/{questionnaire_id}")
def generate_answers(questionnaire_id: int, db: Session = Depends(get_db)):

    questions = db.query(Question).filter(
        Question.questionnaire_id == questionnaire_id
    ).all()

    if not questions:
        return {"error": "No questions found"}

    index, chunks, metadata, vectorizer = load_reference_docs()

    results = []

    for q in questions:

        query_vector = vectorizer.transform([q.text]).toarray().astype("float32")
        D, I = index.search(query_vector, k=1)

        best_chunk = chunks[I[0][0]]
        source = metadata[I[0][0]]
        distance = D[0][0]

        if distance > 1.5:
            answer_text = "Not found in references."
            citation = "None"
            confidence = "Low"
        else:
            # Split chunk into sentences
            sentences = best_chunk.split(".")
            best_sentence = ""

            for sentence in sentences:
                if any(word.lower() in sentence.lower() for word in q.text.split()):
                    best_sentence = sentence.strip()
                    break

            if best_sentence:
                answer_text = best_sentence + "."
            else:
                answer_text = best_chunk.strip()
            citation = source
            confidence = "High" if distance < 0.5 else "Medium"

        answer = Answer(
            question_id=q.id,
            answer=answer_text,
            citation=citation,
            confidence=confidence
        )

        db.add(answer)

        results.append({
            "question": q.text,
            "answer": answer_text,
            "citation": citation,
            "confidence": confidence
        })

    db.commit()

    return {"results": results}

@app.get("/answers/{questionnaire_id}")
def get_answers(questionnaire_id: int, db: Session = Depends(get_db)):

    questions = db.query(Question).filter(
        Question.questionnaire_id == questionnaire_id
    ).all()

    results = []

    for q in questions:
        answer = db.query(Answer).filter(
            Answer.question_id == q.id
        ).first()

        results.append({
            "question_id": q.id,
            "question": q.text,
            "answer": answer.answer if answer else None,
            "citation": answer.citation if answer else None,
            "confidence": answer.confidence if answer else None
        })

    total = len(results)
    answered = sum(
        1 for r in results
        if r["answer"] and r["answer"] != "Not found in references."
    )

    return {
        "summary": {
            "total_questions": total,
            "answered": answered,
            "not_found": total - answered
        },
        "results": results
    }



@app.put("/edit_answer/{answer_id}")
def edit_answer(answer_id: int, new_answer: str, db: Session = Depends(get_db)):

    answer = db.query(Answer).filter(Answer.id == answer_id).first()

    if not answer:
        return {"error": "Answer not found"}

    answer.answer = new_answer
    db.commit()

    return {"message": "Answer updated successfully"}

@app.get("/export/{questionnaire_id}")
def export_document(questionnaire_id: int, db: Session = Depends(get_db)):

    questions = db.query(Question).filter(
        Question.questionnaire_id == questionnaire_id
    ).all()

    doc = Document()

    for q in questions:
        answer = db.query(Answer).filter(
            Answer.question_id == q.id
        ).first()

        doc.add_paragraph(f"Question: {q.text}")

        if answer:
            doc.add_paragraph(f"Answer: {answer.answer}")
            doc.add_paragraph(f"Citation: {answer.citation}")
            doc.add_paragraph(f"Confidence: {answer.confidence}")
        else:
            doc.add_paragraph("Answer: Not generated")

        doc.add_paragraph("")

    file_path = f"exports/questionnaire_{questionnaire_id}.docx"
    doc.save(file_path)

    return FileResponse(file_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=f"questionnaire_{questionnaire_id}.docx")