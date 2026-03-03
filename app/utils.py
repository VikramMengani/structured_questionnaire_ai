from docx import Document

def export_answers(questions_answers, output_path):
    doc = Document()

    for qa in questions_answers:
        doc.add_paragraph(f"Question: {qa['question']}")
        doc.add_paragraph(f"Answer: {qa['answer']}")
        doc.add_paragraph(f"Citation: {qa['citation']}")
        doc.add_paragraph("")

    doc.save(output_path)